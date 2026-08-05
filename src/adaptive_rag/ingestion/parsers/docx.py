"""Extractor de texto de documentos DOCX (OOXML)."""

from __future__ import annotations

from io import BytesIO
from typing import Any, Final
from zipfile import ZipFile

from docx import Document
from docx.opc.exceptions import PackageNotFoundError

from adaptive_rag.ingestion.types import (
    MAX_EXTRACTED_TEXT_CHARS,
    IngestionPipelineError,
    ParsedDocument,
    normalize_text,
)

PARSER_NAME = "docx_text"
PARSER_VERSION = "docx_text_v1"

# Cap on the summed decompressed size of zip members, checked from the
# archive directory before parsing so bomb payloads never reach lxml.
MAX_DOCX_DECOMPRESSED_BYTES: Final[int] = 64 * 1024 * 1024


class DocxTextParser:
    """Parser determinista de parrafos y celdas de tabla via python-docx."""

    parser_version = PARSER_VERSION

    def parse(self, content: bytes) -> ParsedDocument:
        if not content:
            raise IngestionPipelineError("DOCX extraction produced no text")
        _reject_oversized_package(content)
        try:
            document = Document(BytesIO(content))
        except PackageNotFoundError as exc:
            raise IngestionPipelineError(
                "DOCX content is not a readable DOCX package"
            ) from exc
        except Exception as exc:
            raise IngestionPipelineError(
                "DOCX content is not a readable DOCX package"
            ) from exc

        parts: list[str] = []
        extracted_chars = 0
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if text:
                extracted_chars += len(text)
                parts.append(text)
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    extracted_chars += sum(len(cell) for cell in cells)
                    parts.append(" | ".join(cells))
        if extracted_chars > MAX_EXTRACTED_TEXT_CHARS:
            raise IngestionPipelineError("DOCX extracted text exceeds max size")

        normalized = normalize_text("\n\n".join(parts))
        if not normalized:
            raise IngestionPipelineError("DOCX extraction produced no text")

        metadata: dict[str, Any] = {
            "paragraph_count": len(document.paragraphs),
            "table_count": len(document.tables),
        }
        return ParsedDocument(
            normalized_text=normalized,
            parser_metadata={
                "parser": PARSER_NAME,
                "parser_version": self.parser_version,
                "source_type": "docx",
            },
            extraction_metadata=metadata,
        )


def _reject_oversized_package(content: bytes) -> None:
    """Rechaza paquetes cuyos miembros descomprimen por encima del cap.

    Lee solo el directorio del zip (sin descomprimir); si el archivo no es un
    zip valido, se delega el error a python-docx para un mensaje consistente.
    """

    try:
        with ZipFile(BytesIO(content)) as archive:
            total = sum(info.file_size for info in archive.infolist())
    except Exception:
        return
    if total > MAX_DOCX_DECOMPRESSED_BYTES:
        raise IngestionPipelineError(
            "DOCX package exceeds max decompressed size"
        )
